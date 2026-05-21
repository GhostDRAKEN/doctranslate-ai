import shutil
from base64 import b64decode
import json

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.testclient import TestClient

from app.main import app
from app.services import translation_service
from app.services.docx.docx_generator import generate_docx
from app.services.extraction_service import extract_document_intermediate
from app.services.mock_translation_provider import MockTranslationProvider
from app.services.storage_service import (
    get_docx_result_path,
    get_images_directory,
    get_intermediate_path,
    get_storage_root,
    save_source_pdf,
)


def _cleanup_document(document_id: str) -> None:
    document_dir = get_storage_root() / document_id
    if document_dir.exists():
        shutil.rmtree(document_dir)


def _sample_pdf_bytes() -> bytes:
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Service Agreement", fontsize=18)
    page.insert_text(
        (72, 120),
        "This agreement defines the responsibilities of each party.",
        fontsize=11,
    )
    content = pdf.tobytes()
    pdf.close()
    return content


def _pdf_with_image() -> bytes:
    png_bytes = b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
        "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    )
    pdf = fitz.open()
    page = pdf.new_page(width=595, height=842)
    page.insert_text((72, 80), "Image example paragraph.", fontsize=11)
    page.insert_image(fitz.Rect(72, 140, 172, 240), stream=png_bytes)
    content = pdf.tobytes()
    pdf.close()
    return content


def _processed_document_id(client: TestClient) -> str:
    translation_service.build_translation_provider = lambda: MockTranslationProvider()
    upload_response = client.post(
        "/api/documents/upload",
        files={
            "file": (
                "sample.pdf",
                _sample_pdf_bytes(),
                "application/pdf",
            )
        },
    )
    assert upload_response.status_code == 201
    document_id = str(upload_response.json()["document_id"])

    process_response = client.post(
        f"/api/documents/{document_id}/process",
        json={"target_language": "fr", "glossary": []},
    )
    assert process_response.status_code == 202
    return document_id


def _intermediate_with_image(document_id: str, image_path: str) -> dict:
    return {
        "document_id": document_id,
        "source_language": "en",
        "target_language": "fr",
        "domain": "general",
        "metadata": {
            "filename": "source.pdf",
            "page_count": 1,
            "file_size_mb": 0.01,
            "created_at": "2026-05-20T10:00:00Z",
        },
        "mvp_limits": {
            "max_pages": 10,
            "max_file_size_mb": 10,
            "digital_pdf_only": True,
            "requires_selectable_text": True,
        },
        "glossary": [],
        "pages": [
            {
                "page_number": 1,
                "width": 595,
                "height": 842,
                "blocks": [
                    {
                        "id": "block_001",
                        "page_number": 1,
                        "type": "paragraph",
                        "source_text": "Image paragraph",
                        "translated_text": "[FR MOCK] Image paragraph",
                        "bbox": [72, 80, 300, 110],
                        "style": {
                            "font": "Helvetica",
                            "size": 11,
                            "bold": False,
                            "italic": False,
                            "color": "#000000",
                            "alignment": "left",
                        },
                        "reading_order": 1,
                        "status": "translated",
                        "warnings": [],
                    },
                    {
                        "id": "block_002",
                        "page_number": 1,
                        "type": "image",
                        "source_text": "",
                        "translated_text": "",
                        "bbox": [72, 140, 172, 240],
                        "style": {
                            "font": None,
                            "size": None,
                            "bold": False,
                            "italic": False,
                            "color": None,
                            "alignment": "center",
                        },
                        "reading_order": 2,
                        "status": "skipped",
                        "warnings": [],
                        "image_path": image_path,
                        "has_possible_text": False,
                    },
                ],
            }
        ],
        "warnings": [],
    }


def _intermediate_with_repeating_footer(document_id: str) -> dict:
    payload = _intermediate_with_image(document_id, "")
    payload["pages"][0]["blocks"] = [
        {
            "id": "block_001",
            "page_number": 1,
            "source_page": 1,
            "type": "paragraph",
            "role": "body",
            "confidence_score": 0.72,
            "source_text": "Main content",
            "translated_text": "[FR MOCK] Main content",
            "bbox": [72, 100, 300, 130],
            "style": {
                "font": "Helvetica",
                "size": 11,
                "bold": False,
                "italic": False,
                "color": "#000000",
                "alignment": "left",
            },
            "reading_order": 1,
            "status": "translated",
            "warnings": [],
        },
        {
            "id": "block_002",
            "page_number": 1,
            "source_page": 1,
            "type": "footer",
            "role": "repeating_footer",
            "confidence_score": 0.9,
            "source_text": "Company Confidential",
            "translated_text": "[FR MOCK] Company Confidential",
            "bbox": [72, 820, 220, 832],
            "style": {
                "font": "Helvetica",
                "size": 8,
                "bold": False,
                "italic": False,
                "color": "#000000",
                "alignment": "left",
            },
            "reading_order": 2,
            "status": "translated",
            "warnings": ["repeating_page_artifact"],
        },
        {
            "id": "block_003",
            "page_number": 1,
            "source_page": 1,
            "type": "list_item",
            "role": "list",
            "confidence_score": 0.82,
            "source_text": "- First item",
            "translated_text": "[FR MOCK] - First item",
            "bbox": [72, 150, 220, 165],
            "style": {
                "font": "Helvetica",
                "size": 11,
                "bold": False,
                "italic": False,
                "color": "#000000",
                "alignment": "left",
            },
            "reading_order": 3,
            "status": "translated",
            "warnings": [],
        },
    ]
    return payload


def _intermediate_with_styles(document_id: str) -> dict:
    payload = _intermediate_with_repeating_footer(document_id)
    payload["pages"][0]["blocks"] = [
        {
            "id": "block_001",
            "page_number": 1,
            "type": "title",
            "source_text": "Main Title",
            "translated_text": "Titre principal",
            "bbox": [72, 80, 300, 110],
            "style": {
                "font": "Helvetica",
                "size": 20,
                "bold": True,
                "italic": False,
                "color": "#000000",
                "alignment": "center",
            },
            "reading_order": 1,
            "status": "translated",
            "warnings": [],
        },
        {
            "id": "block_002",
            "page_number": 1,
            "type": "list_item",
            "source_text": "- First item",
            "translated_text": "[FR MOCK] - First item",
            "bbox": [72, 130, 300, 150],
            "style": {
                "font": "Helvetica",
                "size": 11,
                "bold": False,
                "italic": True,
                "color": "#000000",
                "alignment": "left",
            },
            "reading_order": 2,
            "status": "translated",
            "warnings": [],
        },
        {
            "id": "block_003",
            "page_number": 1,
            "type": "footer",
            "role": "page_footer",
            "source_text": "Footer note",
            "translated_text": "Note de bas de page",
            "bbox": [72, 820, 220, 832],
            "style": {
                "font": "Helvetica",
                "size": 8,
                "bold": False,
                "italic": False,
                "color": "#000000",
                "alignment": "right",
            },
            "reading_order": 3,
            "status": "translated",
            "warnings": [],
        },
    ]
    return payload


def test_generate_docx_valid_document() -> None:
    client = TestClient(app)
    document_id = _processed_document_id(client)

    response = client.post(f"/api/documents/{document_id}/generate-docx")

    assert response.status_code == 200
    assert response.json() == {
        "document_id": document_id,
        "status": "docx_generated",
        "download_url": f"/api/documents/{document_id}/download/docx",
    }

    docx_path = get_docx_result_path(document_id)
    assert docx_path.is_file()
    docx = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    assert "[FR MOCK] Service Agreement" in text

    _cleanup_document(document_id)


def test_generate_docx_unknown_document_returns_not_found() -> None:
    client = TestClient(app)

    response = client.post("/api/documents/doc_missing/generate-docx")

    assert response.status_code == 404
    assert response.json()["error"]["code"] == "DOCUMENT_NOT_FOUND"


def test_download_docx_returns_file() -> None:
    client = TestClient(app)
    document_id = _processed_document_id(client)
    generate_response = client.post(f"/api/documents/{document_id}/generate-docx")
    assert generate_response.status_code == 200

    response = client.get(f"/api/documents/{document_id}/download/docx")

    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )
    assert response.content.startswith(b"PK")

    _cleanup_document(document_id)


def test_generate_docx_inserts_simple_image() -> None:
    document_id = "doc_docx_image"
    _cleanup_document(document_id)
    save_source_pdf(document_id, _sample_pdf_bytes())
    images_dir = get_images_directory(document_id)
    images_dir.mkdir(parents=True, exist_ok=True)
    image_path = images_dir / "image_001.png"
    image_path.write_bytes(
        b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
            "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    get_intermediate_path(document_id).write_text(
        json.dumps(_intermediate_with_image(document_id, str(image_path))),
        encoding="utf-8",
    )

    docx_path = generate_docx(document_id)
    docx = Document(docx_path)

    assert len(docx.inline_shapes) == 1

    _cleanup_document(document_id)


def test_generate_docx_contains_image_from_extracted_pdf() -> None:
    document_id = "doc_docx_extracted_image"
    _cleanup_document(document_id)
    save_source_pdf(document_id, _pdf_with_image())
    extract_document_intermediate(document_id)

    docx_path = generate_docx(document_id)
    docx = Document(docx_path)

    assert len(docx.inline_shapes) >= 1

    _cleanup_document(document_id)


def test_generate_docx_adds_note_when_image_missing() -> None:
    document_id = "doc_docx_missing_image"
    _cleanup_document(document_id)
    save_source_pdf(document_id, _sample_pdf_bytes())
    get_intermediate_path(document_id).write_text(
        json.dumps(
            _intermediate_with_image(
                document_id,
                str(get_images_directory(document_id) / "missing.png"),
            )
        ),
        encoding="utf-8",
    )

    docx_path = generate_docx(document_id)
    docx = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in docx.paragraphs)

    assert "[Image non inseree]" in text

    _cleanup_document(document_id)


def test_generate_docx_skips_repeating_footer_and_keeps_list() -> None:
    document_id = "doc_docx_repeating_footer"
    _cleanup_document(document_id)
    save_source_pdf(document_id, _sample_pdf_bytes())
    get_intermediate_path(document_id).write_text(
        json.dumps(_intermediate_with_repeating_footer(document_id)),
        encoding="utf-8",
    )

    docx_path = generate_docx(document_id)
    docx = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in docx.paragraphs)

    assert "[FR MOCK] Main content" in text
    assert "Company Confidential" not in text
    assert "First item" in text

    _cleanup_document(document_id)


def test_generate_docx_applies_text_styles() -> None:
    document_id = "doc_docx_styles"
    _cleanup_document(document_id)
    save_source_pdf(document_id, _sample_pdf_bytes())
    get_intermediate_path(document_id).write_text(
        json.dumps(_intermediate_with_styles(document_id)),
        encoding="utf-8",
    )

    docx_path = generate_docx(document_id)
    docx = Document(docx_path)
    paragraphs = [paragraph for paragraph in docx.paragraphs if paragraph.text]

    heading = next(paragraph for paragraph in paragraphs if "Titre principal" in paragraph.text)
    list_item = next(paragraph for paragraph in paragraphs if "First item" in paragraph.text)
    footer = next(paragraph for paragraph in paragraphs if "Note de bas de page" in paragraph.text)

    assert heading.style.name == "Heading 1"
    assert heading.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert heading.runs[0].bold is True
    assert list_item.style.name == "List Bullet"
    assert list_item.runs[0].italic is True
    assert footer.runs[0].font.size.pt == 8
    assert footer.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    _cleanup_document(document_id)
